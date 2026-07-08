"""
templates/modern.py
A clean, single-column "modern" resume template with a colored name header.
build_resume(data, output_path) takes a dict of resume data and produces a .docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)  # navy


def _add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_COLOR
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    # bottom border line
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_resume(data: dict, output_path: str) -> str:
    """
    data keys expected:
      name, title, email, phone, linkedin, github, location,
      summary, skills (list), education (list of dicts: degree, school, year, score),
      experience (list of dicts: role, company, duration, points[list]),
      projects (list of dicts: title, description, tech, link),
      certifications (list)
    """
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Name header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = ACCENT_COLOR

    if data.get("title"):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(data["title"])
        title_run.font.size = Pt(12)
        title_run.italic = True

    # Contact line
    contact_bits = [b for b in [
        data.get("email"), data.get("phone"), data.get("location"),
        data.get("linkedin"), data.get("github")
    ] if b]
    if contact_bits:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact_p.add_run(" | ".join(contact_bits))
        run.font.size = Pt(10)

    # Summary
    if data.get("summary"):
        _add_heading(doc, "Professional Summary")
        doc.add_paragraph(data["summary"])

    # Skills
    if data.get("skills"):
        _add_heading(doc, "Skills")
        doc.add_paragraph(", ".join(data["skills"]))

    # Experience
    if data.get("experience"):
        _add_heading(doc, "Experience")
        for exp in data["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('role','')} — {exp.get('company','')}")
            run.bold = True
            if exp.get("duration"):
                run2 = p.add_run(f"  ({exp['duration']})")
                run2.italic = True
                run2.font.size = Pt(10)
            for point in exp.get("points", []):
                bullet = doc.add_paragraph(point, style="List Bullet")

    # Projects
    if data.get("projects"):
        _add_heading(doc, "Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("title", ""))
            run.bold = True
            if proj.get("link"):
                p.add_run(f"  ({proj['link']})").font.size = Pt(9)
            if proj.get("description"):
                doc.add_paragraph(proj["description"])
            if proj.get("tech"):
                tech_p = doc.add_paragraph()
                tech_run = tech_p.add_run(f"Tech: {proj['tech']}")
                tech_run.italic = True
                tech_run.font.size = Pt(9)

    # Education
    if data.get("education"):
        _add_heading(doc, "Education")
        for edu in data["education"]:
            line = f"{edu.get('degree','')} — {edu.get('school','')}"
            if edu.get("year"):
                line += f" ({edu['year']})"
            if edu.get("score"):
                line += f", {edu['score']}"
            doc.add_paragraph(line)

    # Certifications
    if data.get("certifications"):
        _add_heading(doc, "Certifications")
        for cert in data["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    doc.save(output_path)
    return output_path