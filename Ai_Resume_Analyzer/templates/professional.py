"""
templates/professional.py
A traditional, ATS-friendly "professional" resume template: left-aligned,
minimal color, single font — optimized to parse cleanly in ATS systems.
build_resume(data, output_path) takes a dict of resume data and produces a .docx.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    # simple underline via border-less approach: add a thin rule paragraph
    rule = doc.add_paragraph()
    rule_run = rule.add_run("_" * 90)
    rule_run.font.size = Pt(6)
    rule.paragraph_format.space_after = Pt(2)


def build_resume(data: dict, output_path: str) -> str:
    """
    Same `data` schema as templates/modern.py — this template just renders
    it in a plainer, more traditional single-column ATS-safe layout.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Name header (left aligned, bold, no color — ATS safe)
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(data.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(18)

    if data.get("title"):
        title_p = doc.add_paragraph()
        title_p.add_run(data["title"]).italic = True

    contact_bits = [b for b in [
        data.get("email"), data.get("phone"), data.get("location"),
        data.get("linkedin"), data.get("github")
    ] if b]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))

    if data.get("summary"):
        _add_heading(doc, "Summary")
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        _add_heading(doc, "Skills")
        doc.add_paragraph(", ".join(data["skills"]))

    if data.get("experience"):
        _add_heading(doc, "Experience")
        for exp in data["experience"]:
            line = f"{exp.get('role','')}, {exp.get('company','')}"
            if exp.get("duration"):
                line += f" — {exp['duration']}"
            p = doc.add_paragraph()
            p.add_run(line).bold = True
            for point in exp.get("points", []):
                doc.add_paragraph(point, style="List Bullet")

    if data.get("projects"):
        _add_heading(doc, "Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            p.add_run(proj.get("title", "")).bold = True
            if proj.get("description"):
                doc.add_paragraph(proj["description"])
            if proj.get("tech"):
                doc.add_paragraph(f"Tech Used: {proj['tech']}")
            if proj.get("link"):
                doc.add_paragraph(f"Link: {proj['link']}")

    if data.get("education"):
        _add_heading(doc, "Education")
        for edu in data["education"]:
            line = f"{edu.get('degree','')}, {edu.get('school','')}"
            if edu.get("year"):
                line += f" — {edu['year']}"
            if edu.get("score"):
                line += f" ({edu['score']})"
            doc.add_paragraph(line)

    if data.get("certifications"):
        _add_heading(doc, "Certifications")
        for cert in data["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    doc.save(output_path)
    return output_path